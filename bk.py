import os
import json
import logging
import re
import time
import smtplib
from datetime import datetime, timezone, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.encoders import encode_base64
from typing import List, Dict, Set, Any
from dataclasses import dataclass
import requests
from docx import Document
from dotenv import load_dotenv

# Native Google GenAI SDK
from google import genai
from google.genai import types

load_dotenv()

logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - [%(levelname)s] - %(message)s'
)

@dataclass
class TenderOpportunity:
    id: str
    title: str
    company: str
    country: str
    description: str
    apply_url: str
    source_portal: str
    posted_at: datetime  # Tracks exactly when the opportunity was published

class StudioAturiProcurementHunter:
    def __init__(self):
        gemini_key = os.getenv("GEMINI_API_KEY")
        if not gemini_key:
            logging.error("CRITICAL: GEMINI_API_KEY environment variable is missing!")
            
        self.ai_client = genai.Client(api_key=gemini_key)
        self.target_email = "antonynduhiu26@gmail.com"
        self.history_file = "processed_tenders.json"
        
        # Core Asset Template Paths
        self.fob_template = "Studio_Aturi_Form_of_Bid.docx"
        self.financial_template = "Studio_Aturi_Financial_Proposal.docx"
        self.nda_template = "Studio_Aturi_Mutual_NDA.docx"
        
        # Highly targeted localized country profile configs
        self.country_profiles = {
            "Kenya": {
                "keywords": ["Corporate Brand Strategy", "Brand Mergers & Acquisitions", "Fintech Identity", "Financial Services Branding", "FMCG Packaging Design", "Product SKU Design", "Consumer Insights", "Market Discovery", "Corporate Identity Guidelines", "Brand Manual", "Value Proposition Development", "Stakeholder Perception Survey"],
                "urls": ["https://tenderflow.co.ke", "https://developmentaid.org", "https://safaricom.co.ke/about/suppliers", "https://eastafricatenders.com"]
            },
            "Uganda": {
                "keywords": ["Brand Audit & Diagnostic", "Corporate Profile", "Visual Identity Review", "Customer Experience Strategy", "CX Strategy", "Audience Profiling", "NGO Campaign Branding", "Commercial Product Packaging"],
                "urls": ["https://tenders.unp.me", "https://kazitenders.co.ug", "https://ungm.org"]
            },
            "Tanzania": {
                "keywords": ["Corporate Image", "Reputation Management", "Brand Architecture Development", "Strategic Positioning", "Private Sector Packaging Design", "Insurance & Pensions Private Fund Communication", "Rollout Management", "Brand Launch Activation"],
                "urls": ["https://tanzaniatenders.com", "https://zoomtanzania.com"]
            },
            "Rwanda": {
                "keywords": ["Digital Transformation Branding", "Tech Brand Strategy", "Service Design", "Product Innovation", "Design Thinking Framework", "Human-Centered Design Research", "HCD Research", "Brand Advisory Services", "Organizational Rebranding", "Perception Mapping"],
                "urls": ["https://jobinrwanda.com/tenders", "https://psf.org.rw"]
            },
            "Congo": {
                "keywords": ["Refonte de l'Identité Visuelle", "Identité Corporative", "Stratégie de Marque", "Conception d’Emballage FMCG", "Communication de Changement de Culture", "Creative Direction", "Brand Asset Management"],
                "urls": ["https://mediacongo.net", "https://congovirtuel.com"]
            },
            "Dubai": {
                "keywords": ["Brand Positioning", "Naming Architecture", "Brand Advisory", "Strategy Pivot", "Luxury Packaging Design", "Advanced Marketing", "Creative Strategy", "Customer Experience Strategy", "Journey Mapping", "Value Proposition Development", "Fintech Identity"],
                "urls": ["https://tenderuae.com", "https://tejari.com"]
            },
            "Ethiopia": {
                "keywords": ["Corporate Rebranding and Strategy", "Brand Architecture", "Identity Manual", "Consumer Insights", "Audience Profiling", "Export Product SKU Design", "Strategic Positioning"],
                "urls": ["https://tenders.2merkato.com", "https://thereporterethiopia.com"]
            }
        }
        
        self.processed_tender_ids = self._load_history()

    def _load_history(self) -> Set[str]:
        if os.path.exists(self.history_file):
            try:
                with open(self.history_file, 'r') as f:
                    return set(json.load(f))
            except Exception as e:
                logging.error(f"Error loading tracking history data: {str(e)}")
        return set()

    def _save_history(self):
        try:
            with open(self.history_file, 'w') as f:
                json.dump(list(self.processed_tender_ids), f)
        except Exception as e:
            logging.error(f"Error persisting tracking history data: {str(e)}")

    def scrape_all_opportunities(self) -> List[TenderOpportunity]:
        """Scans private commercial portals and simulates incoming data packages

        to pass clean structures to downstream AI generation layers.
        """
        logging.info("Initializing regional B2B and international NGO scraping matrix...")
        found_tenders = []
        now = datetime.now(timezone.utc)

        # Baseline Web Request Pipeline simulation to catch exposed private listings
        for country, metadata in self.country_profiles.items():
            for target_url in metadata["urls"]:
                try:
                    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
                    response = requests.get(target_url, timeout=10, headers=headers)
                    logging.info(f"Connected to portal {target_url} for territory: {country} (Status: {response.status_code})")
                except Exception as e:
                    logging.warning(f"Bypassing active scrape block on downstream endpoint {target_url}: {e}")

        # Mock Pipeline simulating realistic incoming metadata feeds with explicit posting timestamps
        mock_pipeline_feed = [
            {
                "id": "opp_ke_7721",
                "title": "RFP for Corporate Rebranding & Fintech Identity Architecture",
                "company": "Equator Regional Microfinance Bank PLC",
                "country": "Kenya",
                "source_portal": "TenderFlow Kenya",
                "apply_url": "https://tenderflow.co.ke/view/eq-rebrand-7721",
                "description": "Equator Microfinance Bank requires a creative partner for an institutional strategy pivot and Fintech Identity setup. Deliverables include stakeholder perception surveys, brand strategy positioning, and an updated digital-first brand manual. Contact: procurement@equatormfbank.com",
                "posted_at": now - timedelta(hours=2) # 2 hours old (Valid: Should be processed)
            },
            {
                "id": "opp_ae_9910",
                "title": "Expression of Interest: Luxury Packaging Design and Product Naming Architecture",
                "company": "Al-Mansoori Consumer Goods Holding",
                "country": "Dubai",
                "source_portal": "Tejari Sourcing Network",
                "apply_url": "https://tejari.com/sourcing/amch-luxury-packaging",
                "description": "Seeking an agency to deliver luxury packaging design structures and matching naming architectures for a new premium confectionery lineup. Submit credentials and baseline methodology to portal link.",
                "posted_at": now - timedelta(hours=36) # 36 hours old (Stale: Should be ignored)
            }
        ]

        for entry in mock_pipeline_feed:
            if entry["id"] in self.processed_tender_ids:
                continue
                
            # --- 24-HOUR TIME GATEKEEPER CONSTRAINT ---
            time_difference = now - entry["posted_at"]
            if time_difference > timedelta(hours=24):
                logging.info(f"Skipping stale opportunity {entry['id']} ('{entry['title']}'). Posted {time_difference.total_seconds() / 3600:.1f} hours ago.")
                continue
            # ------------------------------------------

            desc_lower = entry["description"].lower()
            title_lower = entry["title"].lower()
            country_rules = self.country_profiles[entry["country"]]["keywords"]
            
            keyword_match = any(k.lower() in desc_lower or k.lower() in title_lower for k in country_rules)
            
            if keyword_match:
                found_tenders.append(TenderOpportunity(
                    id=entry["id"], title=entry["title"], company=entry["company"],
                    country=entry["country"], description=entry["description"],
                    apply_url=entry["apply_url"], source_portal=entry["source_portal"], 
                    posted_at=entry["posted_at"]
                ))

        return found_tenders

    def generate_tender_intelligence(self, tender: TenderOpportunity) -> Dict[str, Any]:
        """Utilizes Gemini-2.5-Flash to extract granular metadata and build contextual variables."""
        logging.info(f"Extracting intelligence data framework via Gemini for target: {tender.title}...")
        
        prompt = f"""
        Analyze the following private sector tender/RFP opportunity detail carefully.
        
        Opportunity Title: {tender.title}
        Issuing Entity: {tender.company}
        Country Matrix Source: {tender.country}
        Raw Description Block: {tender.description}
        
        Generate a structured JSON configuration layout containing metadata mappings to customize corporate template blocks.
        The extracted currency should follow localized context (e.g., KES for Kenya, AED for Dubai, USD for global NGOs).
        
        Output your response strictly as a structured JSON object matching this schema exactly:
        {{
           "clean_currency": "e.g., KES",
           "phase1_cost": "Numeric string only e.g., 1,200,000",
           "phase2_cost": "Numeric string only",
           "phase3_cost": "Numeric string only",
           "phase4_cost": "Numeric string only",
           "total_cost": "Numeric string only matching sum of all phases",
           "total_cost_words": "Amount in full words capital letters",
           "rfp_reference_no": "Extracted ref number or generated code if missing",
           "client_address": "Full inferred or generated physical/corporate headquarters address",
           "application_steps_markdown": "Step-by-step application pipeline guide clear text",
           "inferred_requirements_markdown": "Bullet list of estimated technical and corporate requirements based on scope",
           "inferred_details_markdown": "Deep descriptive summary analysis of the strategic commercial opportunity"
        }}
        """

        try:
            response = self.ai_client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.15
                ),
            )
            return json.loads(response.text)
        except Exception as e:
            logging.error(f"Gemini metadata mining execution exception caught: {str(e)}")
            raise e

    def _replace_text_safely(self, paragraph, target_string, replacement_text):
        if target_string in paragraph.text:
            text_accum = "".join([run.text for run in paragraph.runs])
            if target_string in text_accum:
                updated_text = text_accum.replace(target_string, str(replacement_text))
                for i, run in enumerate(paragraph.runs):
                    run.text = updated_text if i == 0 else ""

    def process_and_save_docx_artifacts(self, template_path: str, output_path: str, intel: Dict[str, Any], tender: TenderOpportunity):
        """Walks paragraphs and tables within target template spaces to modify token text placeholders."""
        if not os.path.exists(template_path):
            logging.warning(f"Template system link missing: {template_path}. Skipping dynamic generation layer.")
            return False
            
        doc = Document(template_path)
        
        def run_replacement_loop(paragraphs):
            for p in paragraphs:
                self._replace_text_safely(p, "[Insert Date]", datetime.now().strftime("%B %d, %Y"))
                self._replace_text_safely(p, "[Insert RFP Ref No.]", intel.get("rfp_reference_no", "RFP-REF-GEN"))
                self._replace_text_safely(p, "[Insert Project Name (e.g., Corporate Rebranding & Strategy Advisory)]", tender.title)
                self._replace_text_safely(p, "[Insert Client Company Name & Procurement Committee Address]", f"{tender.company}\n{intel.get('client_address', '')}")
                self._replace_text_safely(p, "[INSERT CLIENT COMPANY NAME]", tender.company)
                self._replace_text_safely(p, "[Insert Client Corporate Name]", tender.company)
                self._replace_text_safely(p, "[Insert Country/Jurisdiction]", tender.country)
                self._replace_text_safely(p, "[Insert Address]", intel.get('client_address', 'Corporate Headquarters'))
                self._replace_text_safely(p, "[Insert Preferred Jurisdiction, e.g., the Republic of Kenya / Dubai (DIFC)]", f"the Republic of {tender.country}" if tender.country != "Dubai" else "Dubai (DIFC)")
                
                self._replace_text_safely(p, "[Insert Currency, e.g., KES / USD / AED]", intel.get("clean_currency", "USD"))
                self._replace_text_safely(p, "[Insert Currency and Total Numeric Amount]", f"{intel.get('clean_currency')} {intel.get('total_cost')}")
                self._replace_text_safely(p, "[Insert Amount in Words]", intel.get("total_cost_words", "SPECIFIED AMOUNT"))
                self._replace_text_safely(p, "[Insert Total]", intel.get("total_cost"))

        run_replacement_loop(doc.paragraphs)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    run_replacement_loop(cell.paragraphs)
                    if "[Insert Amount]" in cell.text:
                        for p in cell.paragraphs:
                            if "Phase 1" in p.text or "Discovery" in cell.text:
                                self._replace_text_safely(p, "[Insert Amount]", intel.get("phase1_cost"))
                            elif "Phase 2" in p.text or "Strategy" in cell.text:
                                self._replace_text_safely(p, "[Insert Amount]", intel.get("phase2_cost"))
                            elif "Phase 3" in p.text or "Visual" in cell.text:
                                self._replace_text_safely(p, "[Insert Amount]", intel.get("phase3_cost"))
                            elif "Phase 4" in p.text or "Rollout" in cell.text:
                                self._replace_text_safely(p, "[Insert Amount]", intel.get("phase4_cost"))

        doc.save(output_path)
        return True

    def create_informational_docx(self, output_path: str, title: str, content_markdown: str):
        doc = Document()
        doc.add_heading(title, level=1)
        p = doc.add_paragraph()
        p.add_run(content_markdown)
        doc.save(output_path)

    def distribute_email_package(self, tender: TenderOpportunity, intel: Dict[str, Any], attached_paths: List[str]):
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = os.getenv("SMTP_SENDER_EMAIL")
        sender_password = os.getenv("SMTP_SENDER_PASSWORD")

        if not sender_email or not sender_password:
            logging.error("Missing SMTP credentials configuration parameters. Email broadcast aborted.")
            return

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = self.target_email
        msg['Subject'] = "bid/tender found"

        email_extract = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', tender.description)
        direct_apply_email = email_extract[0] if email_extract else "N/A (Use online portal links provided below)"

        body_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; color: #333; line-height: 1.5;">
            <h2 style="color: #184C78;">🎯 Automated Private Sector Match Found (Last 24 Hours)</h2>
            <hr/>
            <p><strong>Country Location Source:</strong> <span style="background: #e2e8f0; padding: 2px 6px; border-radius: 4px; font-weight: bold;">{tender.country}</span></p>
            <p><strong>Opportunity Title:</strong> {tender.title}</p>
            <p><strong>Issuing Corporate Enterprise:</strong> {tender.company}</p>
            <p><strong>Scouted via Infrastructure Portal:</strong> {tender.source_portal}</p>
            <p><strong>Posted At:</strong> {tender.posted_at.strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
            
            <h3 style="color: #2b6cb0; margin-top: 20px;">📋 Core Summary Details:</h3>
            <p>{tender.description}</p>
            
            <h3 style="color: #2b6cb0; margin-top: 20px;">🚀 Strategic Execution & Application Steps:</h3>
            <div style="background: #f7fafc; padding: 15px; border-left: 4px solid #3182ce; font-family: monospace; white-space: pre-wrap;">
{intel.get("application_steps_markdown")}
            </div>
            
            <p style="margin-top: 15px;"><strong>Direct Email Application Target:</strong> {direct_apply_email}</p>
            <p><strong>Portal Web Link Address:</strong> <a href="{tender.apply_url}" style="color: #3182ce; font-weight: bold;">{tender.apply_url}</a></p>
            <hr/>
            <p style="font-size: 11px; color: #718096; font-style: italic;">Studio Aturi Automation Engine. Transactional proposal documents have been compiled dynamically and attached below.</p>
        </body>
        </html>
        """
        msg.attach(MIMEText(body_content, 'html'))

        for file_path in attached_paths:
            if os.path.exists(file_path):
                filename = os.path.basename(file_path)
                with open(file_path, 'rb') as f:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(f.read())
                    encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename={filename}')
                    msg.attach(part)

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            server.quit()
            logging.info(f"Complete commercial bid packet successfully dispatched via email to {self.target_email}.")
        except Exception as e:
            logging.error(f"SMTP delivery infrastructure failure exception: {str(e)}")

    def run(self):
        logging.info("Starting automated workflow execution cycle...")
        opportunities = self.scrape_all_opportunities()
        logging.info(f"Verification check completed: Found {len(opportunities)} qualifying recent private sector/NGO briefs.")

        for tender in opportunities:
            try:
                # Intel parameters setup
                intel = self.generate_tender_intelligence(tender)
                
                fob_output = f"Studio_Aturi_Form_of_Bid_{tender.id}.docx"
                fin_output = f"Studio_Aturi_Financial_Proposal_{tender.id}.docx"
                nda_output = f"Studio_Aturi_Mutual_NDA_{tender.id}.docx"
                details_output = f"Opportunity_Details_{tender.id}.docx"
                reqs_output = f"Opportunity_Requirements_{tender.id}.docx"
                
                # Document compilations
                self.process_and_save_docx_artifacts(self.fob_template, fob_output, intel, tender)
                self.process_and_save_docx_artifacts(self.financial_template, fin_output, intel, tender)
                self.process_and_save_docx_artifacts(self.nda_template, nda_output, intel, tender)
                
                self.create_informational_docx(details_output, f"Opportunity Details - {tender.title}", intel.get("inferred_details_markdown", "No details provided."))
                self.create_informational_docx(reqs_output, f"Opportunity Requirements - {tender.title}", intel.get("inferred_requirements_markdown", "No requirements extracted."))
                
                attachment_batch = [fob_output, fin_output, nda_output, details_output, reqs_output]
                
                # Dispatch execution
                self.distribute_email_package(tender, intel, attachment_batch)
                
                # Workspace cleaning
                for asset in attachment_batch:
                    if os.path.exists(asset):
                        os.remove(asset)
                        
                self.processed_tender_ids.add(tender.id)
                self._save_history()
                
            except Exception as e:
                logging.error(f"Process pipeline failed for entry identification code {tender.id}: {str(e)}")

if __name__ == "__main__":
    agent = StudioAturiProcurementHunter()
    agent.run()